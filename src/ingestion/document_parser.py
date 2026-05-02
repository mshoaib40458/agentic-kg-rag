"""
Document Parser — Phase 1
Supports: PDF, DOCX, TXT, Markdown, HTML
Preserves structure, strips boilerplate.
"""

import os
import uuid
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Represents a parsed document with content and metadata."""
    doc_id: str
    filename: str
    file_type: str
    content: str
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = str(uuid.uuid4())


class DocumentParser:
    """
    Multi-format document parser.
    Extracts clean text while preserving document structure.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}

    def parse(self, file_path: str, metadata: Optional[dict] = None) -> ParsedDocument:
        """
        Parse a document file into a ParsedDocument.

        Args:
            file_path: Absolute or relative path to the file.
            metadata: Optional additional metadata to attach.

        Returns:
            ParsedDocument with extracted content.

        Raises:
            ValueError: If file type is not supported.
            FileNotFoundError: If file does not exist.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        doc_id = str(uuid.uuid4())
        base_metadata = {
            "filename": path.name,
            "file_path": str(path.resolve()),
            "file_type": ext,
            "file_size_bytes": path.stat().st_size,
        }
        if metadata:
            base_metadata.update(metadata)

        logger.info(f"Parsing document: {path.name} [{ext}]")

        parser_map = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt": self._parse_txt,
            ".md": self._parse_markdown,
            ".html": self._parse_html,
            ".htm": self._parse_html,
        }

        content = parser_map[ext](path)
        content = self._clean_text(content)

        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            file_type=ext,
            content=content,
            metadata=base_metadata,
        )

    def parse_directory(self, dir_path: str, metadata: Optional[dict] = None) -> list[ParsedDocument]:
        """Parse all supported documents in a directory."""
        documents = []
        dir_path = Path(dir_path)

        for file_path in dir_path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.parse(str(file_path), metadata)
                    documents.append(doc)
                    logger.info(f"✓ Parsed: {file_path.name}")
                except Exception as e:
                    logger.error(f"✗ Failed to parse {file_path.name}: {e}")

        logger.info(f"Parsed {len(documents)} documents from {dir_path}")
        return documents

    # ── Private Parsers ────────────────────────────────────────────

    def _parse_pdf(self, path: Path) -> str:
        try:
            import PyPDF2
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

    def _parse_docx(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading hierarchy
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.split()[-1]
                        prefix = "#" * int(level) if level.isdigit() else "#"
                        text_parts.append(f"{prefix} {para.text}")
                    else:
                        text_parts.append(para.text)
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

    def _parse_txt(self, path: Path) -> str:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _parse_markdown(self, path: Path) -> str:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()  # Keep markdown syntax for structure hints

    def _parse_html(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                soup = BeautifulSoup(f.read(), "lxml")
            # Remove script/style tags
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            raise ImportError("beautifulsoup4/lxml not installed.")

    def _clean_text(self, text: str) -> str:
        """
        Remove boilerplate, normalize whitespace, and clean document text.
        Strips: page numbers, Confidential stamps, Page X of Y, repeated
        header/footer lines, TOC dot-leaders, and horizontal rules. (FR-3)
        """
        import re
        from collections import Counter

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        lines = text.split("\n")

        # ── Pass 1: Remove per-line boilerplate ──────────────
        cleaned_lines = []
        boilerplate_patterns = [
            re.compile(r"^\s*\d+\s*$"),                                         # Standalone page numbers
            re.compile(r"(?i)^\s*page\s+\d+\s+of\s+\d+\s*$"),                 # "Page 3 of 12"
            re.compile(r"(?i)\bconfidential\b.*$"),                             # "Confidential – Do not distribute"
            re.compile(r"(?i)\binternal use only\b"),                           # "Internal Use Only"
            re.compile(r"(?i)\ball rights reserved\b"),                         # "All Rights Reserved"
            re.compile(r"(?i)^\s*proprietary\s*$"),                             # "Proprietary"
            re.compile(r"(?i)^\s*draft\s*[-–]?\s*v?\d*\s*$"),                  # "DRAFT v2"
            re.compile(r"^[\.\s·•\-=]{5,}$"),                                    # Rows of dots/dashes (horizontal rules)
            re.compile(r".{20,}\s*\.{4,}\s*\d+\s*$"),                           # TOC: "Introduction ........ 12"
        ]

        for line in lines:
            stripped = line.strip()
            is_boilerplate = any(p.search(stripped) for p in boilerplate_patterns)
            if not is_boilerplate:
                cleaned_lines.append(line)

        # ── Pass 2: Remove repeated header/footer lines ──────
        # Lines appearing identically 3+ times across the document are likely headers/footers
        line_counts = Counter(l.strip() for l in cleaned_lines if len(l.strip()) > 5)
        repeated = {line for line, count in line_counts.items() if count >= 3}

        final_lines = [
            l for l in cleaned_lines
            if l.strip() not in repeated
        ]

        text = "\n".join(final_lines)

        # ── Pass 3: Normalize whitespace ─────────────────────
        text = re.sub(r"\n{3,}", "\n\n", text)                                   # Max 2 blank lines
        final = [line.rstrip() for line in text.split("\n")]
        return "\n".join(final).strip()

